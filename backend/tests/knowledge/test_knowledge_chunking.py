from __future__ import annotations

from pathlib import Path

from docx import Document

from app.services.knowledge_ingestion import chunk_interview_notes, extract_docx_paragraphs


def test_extract_docx_paragraphs_reads_non_empty_text(tmp_path: Path) -> None:
    docx_path = tmp_path / "bagu.docx"
    document = Document()
    document.add_paragraph("基础篇")
    document.add_paragraph("")
    document.add_paragraph("1、 Java语言有哪些特点")
    document.add_paragraph("简单易学，支持多线程。")
    document.save(docx_path)

    paragraphs = extract_docx_paragraphs(docx_path)

    assert paragraphs == ["基础篇", "1、 Java语言有哪些特点", "简单易学，支持多线程。"]


def test_chunk_interview_notes_keeps_section_question_and_overlap() -> None:
    paragraphs = [
        "基础篇",
        "1、 JVM 内存模型是什么",
        "JVM 内存区域包含堆、栈、方法区和程序计数器。" * 18,
        "2、 Redis 为什么快",
        "Redis 主要基于内存、单线程事件循环和高效数据结构。" * 10,
    ]

    chunks = chunk_interview_notes(paragraphs, source_file="10万字总结.docx", target_chars=260, overlap_chars=40)

    assert len(chunks) >= 3
    assert chunks[0].section_path == ["基础篇"]
    assert chunks[0].question == "1、 JVM 内存模型是什么"
    assert "JVM 内存区域" in chunks[0].text
    assert any(chunk.question == "2、 Redis 为什么快" for chunk in chunks)
    assert all(chunk.metadata["content_type"] == "interview_notes" for chunk in chunks)
    assert chunks[1].text[:20] in chunks[0].text[-80:] or chunks[0].text[-20:] in chunks[1].text[:80]


def test_chunk_interview_notes_does_not_treat_numbered_answer_items_as_questions() -> None:
    paragraphs = [
        "基础篇",
        "1、 Java语言有哪些特点",
        "1、简单易学、有丰富的类库",
        "2、面向对象",
        "3、与平台无关性（JVM是Java跨平台使用的根本）",
        "2、面向对象和面向过程的区别",
        "面向过程关注步骤，面向对象关注对象、封装、继承和多态。",
    ]

    chunks = chunk_interview_notes(paragraphs, source_file="10万字总结.docx", target_chars=500, overlap_chars=40)

    assert chunks[0].question == "1、 Java语言有哪些特点"
    assert "1、简单易学" in chunks[0].text
    assert "与平台无关性" in chunks[0].text
    assert chunks[1].question == "2、面向对象和面向过程的区别"
